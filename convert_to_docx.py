#!/usr/bin/env python3
"""
Convert Markdown report to Word document (.docx)
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import markdown
import string

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def format_code_block(paragraph, code_text):
    """Format code blocks with monospace font and background"""
    # Clean code text
    code_text = ''.join(ch for ch in code_text if ch in string.printable and ch not in '\x00\x01\x02\x03\x04\x05\x06\x07\x08\x0b\x0c\x0e\x0f\x10\x11\x12\x13\x14\x15\x16\x17\x18\x19\x1a\x1b\x1c\x1d\x1e\x1f')
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Add light gray background
    run.font.color.rgb = None  # Reset color

def clean_line(line):
    # Remove non-printable/control characters
    return ''.join(ch for ch in line if ch in string.printable and ch not in '\x00\x01\x02\x03\x04\x05\x06\x07\x08\x0b\x0c\x0e\x0f\x10\x11\x12\x13\x14\x15\x16\x17\x18\x19\x1a\x1b\x1c\x1d\x1e\x1f')

def convert_markdown_to_docx(md_file, docx_file):
    """Convert Markdown file to Word document"""
    
    # Read the markdown file
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except UnicodeDecodeError:
        try:
            with open(md_file, 'r', encoding='utf-8-sig') as f:
                md_content = f.read()
        except UnicodeDecodeError:
            with open(md_file, 'r', encoding='latin-1') as f:
                md_content = f.read()
    
    # Create a new Word document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.font.name = 'Calibri'
    
    # Heading 1 style
    h1_style = styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.name = 'Calibri'
    h1_style.font.color.rgb = None
    
    # Heading 2 style
    h2_style = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.name = 'Calibri'
    
    # Heading 3 style
    h3_style = styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.name = 'Calibri'
    
    # Normal text style
    normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.size = Pt(11)
    normal_style.font.name = 'Calibri'
    
    # Split content into lines
    lines = md_content.split('\n')
    
    i = 0
    while i < len(lines):
        line = clean_line(lines[i].strip())
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Title (single #)
        if line.startswith('# '):
            title = line[2:]
            p = doc.add_paragraph(title, style='CustomTitle')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Heading 1 (##)
        if line.startswith('## '):
            heading = line[3:]
            doc.add_paragraph(heading, style='CustomH1')
            i += 1
            continue
        
        # Heading 2 (###)
        if line.startswith('### '):
            heading = line[4:]
            doc.add_paragraph(heading, style='CustomH2')
            i += 1
            continue
        
        # Heading 3 (####)
        if line.startswith('#### '):
            heading = line[5:]
            doc.add_paragraph(heading, style='CustomH3')
            i += 1
            continue
        
        # Horizontal rule
        if line.startswith('---'):
            doc.add_paragraph('_' * 50, style='CustomNormal')
            i += 1
            continue
        
        # Code blocks (```)
        if line.startswith('```'):
            # Find the end of the code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip the closing ```
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(style='CustomNormal')
                format_code_block(p, code_text)
            continue
        
        # Inline code (`)
        if '`' in line:
            # Simple inline code formatting
            line = line.replace('`', '')
            p = doc.add_paragraph(line, style='CustomNormal')
            i += 1
            continue
        
        # Lists
        if line.startswith('- ') or line.startswith('* '):
            item = line[2:]
            p = doc.add_paragraph(f'• {item}', style='CustomNormal')
            i += 1
            continue
        
        # Numbered lists
        if re.match(r'^\d+\. ', line):
            item = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(f'• {item}', style='CustomNormal')
            i += 1
            continue
        
        # Tables (simple handling)
        if line.startswith('|'):
            # Collect table lines
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            # Create table
            if len(table_lines) > 2:  # At least header and separator
                # Parse table
                rows = []
                for table_line in table_lines:
                    if table_line.strip() and not table_line.strip().startswith('|-'):
                        cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                        rows.append(cells)
                
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            table.cell(row_idx, col_idx).text = cell_data
            continue
        
        # Bold text (**)
        if '**' in line:
            line = line.replace('**', '')
            p = doc.add_paragraph(line, style='CustomNormal')
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph(line, style='CustomNormal')
        i += 1
    
    # Save the document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    convert_markdown_to_docx('Report/final_report.md', 'Report/final_report.docx') 